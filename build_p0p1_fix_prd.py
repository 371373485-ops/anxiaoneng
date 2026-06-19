from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path("智能经营诊断_P0P1问题修复_PRD_V1.0.docx")
FONT = "Microsoft YaHei"
NAVY, BLUE, GRAY = "17365D", "2E74B5", "667085"
LIGHT, PALE, WHITE = "F2F4F7", "EAF2F8", "FFFFFF"
RED, AMBER, GREEN, BLACK = "9B1C1C", "9A6700", "176B43", "1A1A1A"


def font(run, size=10.5, bold=False, color=BLACK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), FONT)
    rpr.rFonts.set(qn("w:ascii"), "Arial")
    rpr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold, run.italic = bold, italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, bottom=90, start=120, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tcm = tcpr.first_child_found_in("w:tcMar")
    if tcm is None:
        tcm = OxmlElement("w:tcMar")
        tcpr.append(tcm)
    for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tcm.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tcm.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def geometry(table, widths):
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    if tblw.getparent() is None:
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    ind = tblpr.find(qn("w:tblInd")) or OxmlElement("w:tblInd")
    if ind.getparent() is None:
        tblpr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW")) or OxmlElement("w:tcW")
            if tcw.getparent() is None:
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[i]))
            tcw.set(qn("w:type"), "dxa")
            margins(cell)


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trpr.append(node)


def rule(paragraph, color="D9E2F3", size=6):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, NAVY, 8, 4),
]:
    style = doc.styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size, style.font.bold = Pt(size), True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True


def add_numbering(num_id, abstract_id, fmt, text):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
    nf = OxmlElement("w:numFmt"); nf.set(qn("w:val"), fmt); lvl.append(nf)
    lt = OxmlElement("w:lvlText"); lt.set(qn("w:val"), text); lvl.append(lt)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "720")
    tabs.append(tab); ppr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "720"); ind.set(qn("w:hanging"), "360")
    ppr.append(ind); lvl.append(ppr); abstract.append(lvl); numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    aid = OxmlElement("w:abstractNumId"); aid.set(qn("w:val"), str(abstract_id))
    num.append(aid); numbering.append(num)


add_numbering(31, 31, "bullet", "•")
add_numbering(32, 32, "decimal", "%1.")


def numbered(text, num=31):
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId"); numid.set(qn("w:val"), str(num))
    numpr.append(ilvl); numpr.append(numid); ppr.append(numpr)
    r = p.add_run(text); font(r)
    return p


def para(text, lead=None):
    p = doc.add_paragraph()
    if lead and text.startswith(lead):
        r = p.add_run(lead); font(r, bold=True)
        r = p.add_run(text[len(lead):]); font(r)
    else:
        r = p.add_run(text); font(r)
    return p


def heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def table(headers, rows, widths, aligns=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    geometry(t, widths)
    repeat_header(t.rows[0])
    for i, text in enumerate(headers):
        c = t.rows[0].cells[i]; shade(c, LIGHT); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text)); font(r, 9.2, True, NAVY)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = aligns[i] if aligns else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value)); font(r, 8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def callout(label, text, fill=PALE, color=NAVY):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; geometry(t, [9360])
    c = t.cell(0, 0); shade(c, fill)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " "); font(r, 10.5, True, color)
    r = p.add_run(text); font(r, 10.5, False, color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


# Header/footer
hp = sec.header.paragraphs[0]
r = hp.add_run("产品需求文档  |  智能经营诊断 P0/P1 问题修复")
font(r, 8.5, True, GRAY); rule(hp)
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = fp.add_run("P0/P1修复PRD  |  "); font(r, 8.5, color=GRAY)
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); fp._p.append(fld)

# Cover
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(28); p.paragraph_format.space_after = Pt(5)
r = p.add_run("产品需求文档（修复专项）"); font(r, 12, True, BLUE)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
r = p.add_run("智能经营诊断 P0/P1 问题修复"); font(r, 24, True, NAVY)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(20)
r = p.add_run("以当前已实现版本为基线，面向受控试点和深度经营分析的修复方案"); font(r, 13, color=GRAY)
for k, v in [
    ("版本", "V1.0"),
    ("状态", "产品与研发评审稿"),
    ("基线", "2026年6月18日已实现版本"),
    ("目标版本", "P0受控试点版 / P1经营分析增强版"),
    ("优先级定义", "P0为发布阻断项；P1为试点后增强项"),
    ("编制日期", "2026年6月18日"),
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(k + "："); font(r, 10.5, True, NAVY)
    r = p.add_run(v); font(r)
p = doc.add_paragraph(); rule(p, BLUE, 12)
callout("核心目标", "修复诊断专业度下降、指标对标错误、整改方向误判和真实AI未验收等问题，在不破坏现有规则报告与公开静态版的前提下，形成可信、可测、可试点的智能经营诊断能力。")
doc.add_page_break()

heading("文档修订记录", 1)
table(["版本", "日期", "状态", "说明"], [
    ("V1.0", "2026-06-18", "评审稿", "定义P0阻断问题和P1增强问题的具体修复需求、验收标准与排期")
], [1200, 1800, 1600, 4760],
      [WD_ALIGN_PARAGRAPH.CENTER]*3 + [WD_ALIGN_PARAGRAPH.LEFT])

heading("1. 背景与基线结论", 1)
para("当前版本已经实现智能经营诊断页面、规则诊断、证据追溯、诊断快照、AI结构化接口、追问、反馈、整改任务、审计、权限和Docker部署配置。自动化测试、运行态冒烟、证据抽屉、任务草稿、状态流转和机构越权拦截均已通过验证。")
callout("发布判断", "当前版本适合内部演示和功能验证，但真实AI、诊断专业度、指标对标口径及整改复盘方向尚未达到全面生产发布标准。", "FFF4E5", AMBER)

heading("1.1 问题分级", 2)
table(["级别", "定义", "处理原则"], [
    ("P0", "不修复会导致错误经营结论、错误复盘、权限风险或无法完成真实AI验收", "全部关闭后方可进入受控试点"),
    ("P1", "不阻断试点，但限制分析深度、管理价值或规模化使用", "试点反馈验证后分批上线"),
], [1300, 4560, 3500])

heading("1.2 本期范围", 2)
for text in [
    "P0-01：统一规则诊断内核，恢复原有专业分析深度。",
    "P0-02：建立指标元数据与对标策略，修复绝对金额错误对标。",
    "P0-03：修复整改指标方向和复盘口径。",
    "P0-04：建立真实AI评测、数字校验和发布门槛。",
    "P0-05：机构编码、身份权限和生产部署验证。",
    "P0-06：补齐关键接口、权限、会话、任务及审计自动化测试。",
    "P1-01：多机构智能比较。",
    "P1-02：多月趋势与异常分类。",
    "P1-03：分险种、渠道、赔案和费用下钻。",
    "P1-04：自然语言经营查询与会议摘要。",
    "P1-05：整改工作台增强与运营指标。",
]:
    numbered(text)

heading("2. 产品目标与非目标", 1)
heading("2.1 目标", 2)
for text in [
    "新诊断页面的专业分析能力不得低于原规则报告。",
    "所有指标按业务属性选择正确的对标对象和改善方向。",
    "真实AI输出可被自动验证、可回归、可降级、可审计。",
    "整改复盘能够正确判断指标改善或恶化。",
    "机构权限使用稳定编码，避免名称和Header编码风险。",
    "P1形成从机构级异常发现到业务驱动下钻的分析链路。",
]:
    numbered(text, 32)

heading("2.2 非目标", 2)
for text in [
    "本期不改变原始经营指标计算口径。",
    "本期不允许AI自动修改数据、规则或任务状态。",
    "P0不建设完整工作流平台和消息中心。",
    "P1不替代精算、理赔、财务等专业系统的明细分析。",
]:
    numbered(text)

heading("3. P0需求总览", 1)
table(["需求编号", "需求名称", "业务风险", "发布要求"], [
    ("P0-01", "统一诊断内核", "新页面诊断深度低于原报告", "阻断"),
    ("P0-02", "指标元数据与对标策略", "绝对金额比较误导管理判断", "阻断"),
    ("P0-03", "整改方向与复盘口径", "改善结果可能被判为恶化", "阻断"),
    ("P0-04", "真实AI评测与可信门禁", "代码完成但质量未知", "阻断"),
    ("P0-05", "机构编码与生产权限", "名称编码、别名和越权风险", "阻断"),
    ("P0-06", "关键链路测试补齐", "回归风险不可控", "阻断"),
], [1300, 2350, 3550, 2160],
      [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])

heading("4. P0-01 统一规则诊断内核", 1)
heading("4.1 问题", 2)
para("当前新页面只围绕8项指标构建诊断，风险等级主要按预警数量计算；原规则报告中的多维风险评分、经营模式识别、分层归因、排名考核和短中长期建议没有完整复用。")
heading("4.2 目标方案", 2)
para("建立唯一的结构化诊断服务 buildDiagnosisModel。原规则引擎负责计算，新页面和原报告均消费同一结构化结果，禁止两套风险评分和建议逻辑并行演进。")
table(["模块", "必须输出", "来源"], [
    ("风险等级", "分数、等级、触发因子、规则版本", "原多维风险矩阵"),
    ("关键事实", "当前值、对标值、排名、同比/环比、证据ID", "指标计算与证据层"),
    ("经营模式", "模式名称、触发条件、业务含义", "原经营模式识别"),
    ("归因假设", "推断、置信度、证据ID、缺失数据", "规则诊断"),
    ("核查清单", "核查项、优先级、责任专业", "规则诊断"),
    ("管理建议", "短中长期、动作、责任角色、观察指标", "规则诊断"),
], [1800, 4800, 2760])
heading("4.3 验收标准", 2)
for text in [
    "同一机构、周期、数据版本和规则版本生成完全一致的诊断对象。",
    "新页面至少覆盖原规则报告现有的风险等级、模式识别、三层风险表现、归因、核查和建议。",
    "不存在旧页面显示高风险、新页面显示关注的冲突。",
    "诊断对象包含schemaVersion、dataVersion和ruleVersion。",
    "无预警但存在明显排名或对标异常时，仍能输出关注项。",
]:
    numbered(text)

heading("5. P0-02 指标元数据与对标策略", 1)
heading("5.1 指标元数据模型", 2)
table(["字段", "说明", "示例"], [
    ("metricId", "稳定指标编码", "COR_ACTUAL"),
    ("label/unit", "展示名称与单位", "综合成本率 / %"),
    ("category", "ratio / attainment / amount / productivity / count", "ratio"),
    ("direction", "increase / decrease / target / neutral", "decrease"),
    ("benchmarkStrategy", "weightedOverall / regional / median / plan / prior / peerGroup / none", "weightedOverall"),
    ("trendThreshold", "判断改善、恶化和持平的最小变化", "0.02"),
    ("evidencePrecision", "证据和AI引用精度", "1位百分数"),
], [2200, 3900, 3260])

heading("5.2 默认对标规则", 2)
table(["指标类型", "默认对标", "禁止行为"], [
    ("赔付率/费用率/COR", "全国或责任区加权值；必要时同险种", "使用简单平均"),
    ("保费/利润达成率", "计划目标、同期达成、同区域中位数", "将达成率误当绝对金额"),
    ("人均产能/人均利润", "机构中位数、责任区或同规模组", "直接使用全国总额"),
    ("保费/利润绝对额", "自身计划、同期、同规模机构组；默认不显示全国总额对标", "分公司绝对额与全国合计直接比较"),
    ("人数/费用金额", "预算、同期或同规模组", "仅按数值大小判断好坏"),
], [2300, 4200, 2860])
heading("5.3 页面要求", 2)
for text in [
    "证据抽屉展示“对标对象”和“对标策略”，例如“责任区加权值”或“自身年度计划”。",
    "无合法对标对象时显示“暂不提供对标”，不得显示无业务意义的数字。",
    "对标差异必须按指标方向生成解释，低赔付率为改善，高保费达成率为改善。",
    "所有对标结果由事实层计算，AI只引用结果。",
]:
    numbered(text)
heading("5.4 验收标准", 2)
for text in [
    "经营利润等绝对金额不再与全国合计直接对标。",
    "区域比率使用保费或已赚保费加权，结果与计算引擎一致。",
    "新增指标未配置元数据时禁止进入AI分析，并记录配置缺失日志。",
    "至少覆盖现有全部核心指标的方向和对标策略。",
]:
    numbered(text)

heading("6. P0-03 整改方向与复盘口径", 1)
heading("6.1 修复要求", 2)
para("整改任务不得再通过指标名称是否包含“率”推断方向，必须读取指标元数据 direction。")
table(["direction", "含义", "复盘逻辑"], [
    ("increase", "越高越好", "当前值高于整改前为改善"),
    ("decrease", "越低越好", "当前值低于整改前为改善"),
    ("target", "接近目标值最好", "比较与目标值的距离是否缩小"),
    ("neutral", "仅观察", "只展示变化，不评价改善或恶化"),
], [1800, 3000, 4560])
heading("6.2 复盘输出", 2)
for text in [
    "整改前值、整改后值、变化值和变化比例。",
    "目标值、是否达标和距离目标的变化。",
    "对标值变化及机构排名变化。",
    "明显改善、小幅改善、无明显变化、继续恶化或数据不足。",
    "固定提示：指标变化不等于整改措施与结果之间存在确定因果关系。",
]:
    numbered(text)
heading("6.3 验收场景", 2)
table(["场景", "预期"], [
    ("保费达成率80%升至95%", "判定改善，不得因名称含“率”判为恶化"),
    ("赔付率75%降至68%", "判定改善"),
    ("COR105%降至99%", "判定明显改善"),
    ("目标型指标从90偏离至95，目标100", "判定向目标改善"),
    ("后续周期缺少同一指标", "判定数据不足，不生成结论"),
], [4000, 5360])

heading("7. P0-04 真实AI评测与可信门禁", 1)
heading("7.1 评测环境", 2)
for text in [
    "使用脱敏诊断快照，覆盖高赔付、高费用、规模不足、人效不足、单月数据、分母为零、单位不一致等不少于30个案例。",
    "每个案例包含输入数据、允许结论、禁止结论、必须引用证据和预期建议方向。",
    "固定模型、提示词版本、Schema版本和温度参数，保证可重复评测。",
]:
    numbered(text)
heading("7.2 可信门禁", 2)
table(["指标", "P0发布标准", "阻断条件"], [
    ("数字引用准确率", "≥99%", "任何关键金额、比例或排名与证据不一致"),
    ("无依据结论率", "≤2%", "推断无证据或缺少数据不足声明"),
    ("Schema成功率", "≥98%", "结构错误重试后仍失败"),
    ("关键禁区违规", "0", "编造同比、行业标准、确定性归因"),
    ("首字响应", "P95≤5秒", "P95超过8秒"),
    ("完整响应", "P95≤30秒", "大面积超时或无法取消"),
    ("降级成功率", "100%", "AI失败影响基础诊断"),
], [2880, 2520, 3960])
heading("7.3 产品处理", 2)
for text in [
    "AI输出必须先通过Schema、证据ID和数字校验，再展示给用户。",
    "校验失败自动重试一次；仍失败时显示基础诊断并记录失败原因。",
    "提示词和模型版本进入审计日志，支持按版本比较评测结果。",
    "AI按钮展示服务状态；关闭、限流、超时和安全拦截使用不同文案。",
]:
    numbered(text)

heading("8. P0-05 机构编码、权限与生产部署", 1)
heading("8.1 机构权限模型", 2)
table(["字段", "要求"], [
    ("orgId", "稳定机构编码，不使用中文名称作为权限主键"),
    ("orgName", "展示名称，可调整但不影响权限"),
    ("orgType", "hq / region / branch"),
    ("parentOrgId", "支持总公司、责任区、分公司树形范围"),
    ("allowedOrgIds", "由身份网关或服务端会话注入"),
], [2600, 6760])
heading("8.2 安全要求", 2)
for text in [
    "生产环境只接受可信身份网关注入的标准ASCII Header或服务端Token。",
    "浏览器localStorage中的用户、角色和机构信息仅可用于开发演示，不得作为生产身份依据。",
    "服务端按orgId校验诊断、证据、会话、反馈、任务和复盘全部对象。",
    "管理员审计接口保持单独权限，记录查询人和查询条件。",
    "PostgreSQL和Docker环境必须完成真实启动、迁移、重启和数据持久化验证。",
]:
    numbered(text)
heading("8.3 验收标准", 2)
for text in [
    "分公司用户不能通过修改请求参数访问其他机构数据。",
    "机构改名后历史诊断、任务和审计仍可正确关联。",
    "生产环境未配置proxy或token鉴权时拒绝启动。",
    "应用重启后诊断、任务、反馈和审计数据不丢失。",
]:
    numbered(text)

heading("9. P0-06 测试与发布保障", 1)
table(["测试域", "必须覆盖"], [
    ("规则诊断", "新旧诊断一致性、模式识别、无预警异常、缺失数据"),
    ("指标元数据", "全部核心指标方向、对标策略、未配置指标拦截"),
    ("AI", "成功、Schema错误、数字错误、超时、限流、取消、降级"),
    ("会话", "同机构连续追问、切换机构隔离、历史上限、取消保存"),
    ("权限", "五类角色、跨机构访问、审计接口、伪造Header"),
    ("整改", "创建、确认、执行、完成、关闭、非法跳转、逾期"),
    ("复盘", "四类方向、目标达成、无数据、跨机构/错误周期"),
    ("部署", "SQLite开发、PostgreSQL生产、容器重启、静态公开版"),
], [2200, 7160])
callout("P0完成定义", "全部P0验收用例通过；无Critical/High缺陷；真实AI评测达到门禁；生产鉴权和PostgreSQL部署完成验证；公开静态版仍保持AI关闭和数据不外传。", "EAF7F0", GREEN)

heading("10. P1需求总览", 1)
table(["需求编号", "名称", "核心价值"], [
    ("P1-01", "多机构智能比较", "从单机构诊断升级为差异化管理"),
    ("P1-02", "多月趋势与异常分类", "区分短期波动和方向性变化"),
    ("P1-03", "业务驱动下钻", "从结果指标定位险种、渠道、赔案和费用驱动"),
    ("P1-04", "自然语言查询与会议摘要", "降低管理人员分析和汇报成本"),
    ("P1-05", "整改工作台增强", "形成可运营、可统计的改善闭环"),
], [1500, 3000, 4860])

heading("11. P1-01 多机构智能比较", 1)
heading("11.1 功能", 2)
for text in [
    "支持选择2至5家分公司，默认同责任区或同规模组。",
    "比较经营风险、核心指标、经营模式、趋势和整改状态。",
    "输出共性问题、差异问题、标杆机构和差异化建议。",
    "比较必须使用同周期、同口径和同一计划版本数据。",
]:
    numbered(text)
heading("11.2 交互", 2)
para("诊断工具栏增加“机构比较”。用户选择机构后展示比较矩阵和AI摘要；可从比较结果跳转到单机构证据和整改任务。")
heading("11.3 验收", 2)
for text in [
    "绝对金额按同规模组或自身计划比较，不以全国合计直接排序。",
    "比较周期或口径不一致时阻止生成并说明原因。",
    "所有AI比较结论可追溯到机构和指标证据。",
]:
    numbered(text)

heading("12. P1-02 多月趋势与异常分类", 1)
table(["状态", "判定要求"], [
    ("持续改善", "至少3个连续有效周期，按指标方向连续改善且超过阈值"),
    ("持续恶化", "至少3个连续有效周期，按指标方向连续恶化且超过阈值"),
    ("高位波动", "围绕风险阈值以上波动"),
    ("低位波动", "处于健康区间但月度波动明显"),
    ("单月异常", "当前月显著偏离前期区间，前期相对稳定"),
    ("数据不足", "不足3个有效周期或口径不一致"),
], [2200, 7160])
for text in [
    "趋势阈值由指标元数据配置，不使用统一固定阈值。",
    "趋势证据展示各月数值、变化和口径版本。",
    "AI不得基于单月数据使用“持续”“趋势性”等表述。",
]:
    numbered(text)

heading("13. P1-03 分险种、渠道、赔案和费用下钻", 1)
heading("13.1 数据主题", 2)
table(["主题", "最低数据字段", "典型问题"], [
    ("险种", "险种、保费、已赚保费、赔付、费用、COR", "亏损由哪些险种驱动"),
    ("渠道", "渠道、保费、新单、续保、费用、赔付", "增长是否来自低质量渠道"),
    ("赔案", "赔案号、险种、金额、出险/结案时间、大案标识", "高赔付是否由大案或频度驱动"),
    ("费用", "费用科目、预算、实际、固定/变动属性", "费用偏差来自哪些科目"),
    ("人员", "组织、岗位、前后台、人数、人力成本、产出", "人效问题来自结构还是产出"),
], [1500, 4320, 3540])
heading("13.2 下钻路径", 2)
para("机构异常 → 指标异常 → 业务主题 → 驱动项排序 → 明细核查清单 → 改善建议。")
heading("13.3 约束", 2)
for text in [
    "下钻数据必须标记来源系统、更新时间和口径。",
    "缺少业务明细时只输出核查建议，不输出确定性根因。",
    "涉及赔案和人员明细时按权限脱敏展示。",
]:
    numbered(text)

heading("14. P1-04 自然语言查询与会议摘要", 1)
heading("14.1 查询能力", 2)
for text in [
    "支持“本月高风险机构有哪些”“第三责任区费用率最高的三家机构”等受控查询。",
    "查询先解析为结构化意图和指标编码，再由数据层执行，模型不得直接计算。",
    "返回结果包含筛选条件、口径、周期和证据链接。",
]:
    numbered(text)
heading("14.2 摘要模板", 2)
table(["模板", "输出内容"], [
    ("总公司经营分析会", "全国概况、高风险机构、共性问题、重点整改"),
    ("责任区分析会", "区域对标、机构分层、标杆与落后机构、行动清单"),
    ("分公司专题会", "本机构风险、驱动因素、核查事项、整改进度"),
], [2300, 7060])

heading("15. P1-05 整改工作台增强", 1)
for text in [
    "任务详情页支持编辑记录、状态历史、复盘历史和附件索引。",
    "责任部门和责任人使用组织目录选择器，不使用浏览器prompt。",
    "增加到期提醒、逾期标识、责任区升级和任务筛选。",
    "增加任务按机构、部门、指标、状态和逾期情况统计。",
    "形成建议采纳率、任务完成率、逾期率、目标达成率和指标改善率。",
]:
    numbered(text)

heading("16. 数据与接口改造", 1)
table(["对象", "新增关键字段/接口"], [
    ("metric_metadata", "metric_id、category、direction、benchmark_strategy、trend_threshold、precision"),
    ("organizations", "org_id、org_name、org_type、parent_org_id、status"),
    ("diagnoses", "schema_version、org_id、comparison_scope、diagnosis_score"),
    ("evidence", "benchmark_type、benchmark_label、metric_id、calculation_version"),
    ("remediation_tasks", "direction、target_type、status_history、source_recommendation_id"),
    ("reviews", "rank_change、benchmark_change、distance_to_target_change"),
    ("AI评测", "POST /api/evaluations/run；GET /api/evaluations/{id}"),
    ("多机构比较", "POST /api/comparisons"),
    ("自然语言查询", "POST /api/queries/execute"),
], [2400, 6960])

heading("17. 埋点与运营指标", 1)
table(["指标", "定义"], [
    ("诊断打开率", "查看智能诊断的活跃用户/看板活跃用户"),
    ("AI生成率", "生成AI解读的诊断数/可生成诊断数"),
    ("证据查看率", "点击查看依据的AI或规则结论占比"),
    ("追问率", "发生追问的诊断会话占比"),
    ("建议采纳率", "转为整改任务的建议占比"),
    ("任务完成率", "已完成及关闭任务/已确认任务"),
    ("整改改善率", "复盘结果为明显或小幅改善的任务占比"),
    ("AI失败降级率", "AI失败后成功回退基础诊断的比例"),
], [3000, 6360])

heading("18. 排期与里程碑", 1)
table(["里程碑", "建议周期", "交付内容", "退出条件"], [
    ("M1 P0事实与口径", "1.5周", "统一诊断内核、指标元数据、对标策略、整改方向", "业务口径评审通过"),
    ("M2 P0可信AI与权限", "1.5周", "真实AI评测、机构编码、权限、测试补齐", "P0发布门禁通过"),
    ("M3 受控试点", "2周", "3-5家机构试点、反馈和指标监测", "无高风险缺陷，用户满意率达标"),
    ("M4 P1比较与趋势", "2周", "多机构比较、多月趋势", "功能与口径验收"),
    ("M5 P1下钻与运营", "3-4周", "业务主题下钻、查询、摘要、整改增强", "试点扩围评审通过"),
], [1750, 1550, 3500, 2560])

heading("19. 发布策略", 1)
heading("19.1 P0受控试点", 2)
for text in [
    "默认AI关闭，由管理员按试点机构白名单开启。",
    "保留基础诊断作为始终可用的降级路径。",
    "试点期间每日监控错误、数字校验失败、响应时延和用户反馈。",
    "发现关键数字错误、越权或错误整改复盘时立即关闭AI增强能力。",
]:
    numbered(text)
heading("19.2 P1灰度", 2)
for text in [
    "多机构比较和趋势分析先向总公司及责任区角色开放。",
    "明细下钻按数据主题逐项灰度，并单独完成权限评审。",
    "自然语言查询仅开放白名单意图，不支持任意SQL或自由数据访问。",
]:
    numbered(text)

heading("20. 最终验收清单", 1)
checklist = [
    "新旧页面风险等级和专业诊断逻辑一致。",
    "全部核心指标已配置方向、类型和对标策略。",
    "绝对金额不存在与全国合计直接对标。",
    "整改复盘通过increase/decrease/target/neutral全部用例。",
    "真实AI评测达到数字、Schema、禁区和性能门禁。",
    "机构权限使用orgId并通过跨机构越权测试。",
    "PostgreSQL、容器重启和数据持久化验证完成。",
    "关键API、权限、会话、反馈、任务、复盘和审计测试通过。",
    "公开静态版AI关闭且不发送经营数据。",
    "试点监控指标、回滚方案和责任人已确认。",
]
for item in checklist:
    numbered(item)
callout("建议决策", "先完成全部P0并进入3-5家机构受控试点；P1按“比较与趋势 → 业务下钻 → 查询与整改运营”的顺序推进，不建议并行铺开。", "EAF7F0", GREEN)

doc.core_properties.title = "智能经营诊断P0P1问题修复PRD"
doc.core_properties.subject = "智能经营诊断受控试点与深度分析修复需求"
doc.core_properties.author = "产品团队"
doc.core_properties.keywords = "PRD,P0,P1,经营诊断,财险,AI,整改"
doc.save(OUT)
print(OUT.resolve())
