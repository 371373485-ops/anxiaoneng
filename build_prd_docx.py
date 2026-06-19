from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path


OUT = Path("智能经营诊断与整改闭环_PRD_V1.0.docx")
NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"
GREEN = "176B43"
BLACK = "1A1A1A"
FONT = "Microsoft YaHei"


def set_run_font(run, size=11, bold=False, color=BLACK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, color=MID_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, 9, color=MID_GRAY)


def add_horizontal_rule(paragraph, color=BLUE, size=10):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_num_def(doc, num_id, abstract_id, fmt, text, left=720, hanging=360):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    aid = OxmlElement("w:abstractNumId")
    aid.set(qn("w:val"), str(abstract_id))
    num.append(aid)
    numbering.append(num)


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(numid)
    p_pr.append(num_pr)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.9)
section.bottom_margin = Inches(0.85)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.42)
section.footer_distance = Inches(0.42)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, NAVY, 8, 4),
]:
    st = styles[name]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if "PRD Callout" not in [s.name for s in styles]:
    callout = styles.add_style("PRD Callout", WD_STYLE_TYPE.PARAGRAPH)
    callout.font.name = FONT
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    callout.font.size = Pt(11)
    callout.font.color.rgb = RGBColor.from_string(NAVY)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.15

add_num_def(doc, 21, 21, "bullet", "•", 620, 300)
add_num_def(doc, 22, 22, "decimal", "%1.", 680, 320)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = hp.add_run("产品需求文档 PRD  |  智能经营诊断与整改闭环")
set_run_font(r, 8.5, bold=True, color=MID_GRAY)
add_horizontal_rule(hp, color="D9E2F3", size=4)
add_page_number(section.footer.paragraphs[0])


def add_body(text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(text):
    p = doc.add_paragraph()
    apply_num(p, 21)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(text):
    p = doc.add_paragraph()
    apply_num(p, 22)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_callout(label, text, color=NAVY, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label + " ")
    set_run_font(r1, bold=True, color=color)
    r2 = p.add_run(text)
    set_run_font(r2, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(headers, rows, widths, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(str(text))
        set_run_font(rr, 9.5, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = (alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            rr = p.add_run(str(value))
            set_run_font(rr, 9)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


# Cover / masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(32)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("产品需求文档（PRD）")
set_run_font(r, 12, bold=True, color=BLUE)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("智能经营诊断与整改闭环")
set_run_font(r, 26, bold=True, color=NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(22)
r = p.add_run("分公司经营与效能数据看板 AI 能力升级")
set_run_font(r, 14, color=MID_GRAY)

meta = [
    ("文档版本", "V1.0"),
    ("文档状态", "产品评审稿"),
    ("适用项目", "分公司经营与效能数据看板"),
    ("建设方向", "规则诊断 → AI解释 → 用户追问 → 整改跟踪"),
    ("编制日期", "2026年6月18日"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{label}：")
    set_run_font(r1, 10.5, bold=True, color=NAVY)
    r2 = p.add_run(value)
    set_run_font(r2, 10.5)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
add_horizontal_rule(p, color=BLUE, size=12)

add_callout(
    "产品定位",
    "本产品不是通用聊天机器人，而是基于公司经营数据、预警规则和财险经营知识的智能经营分析助手。",
)

doc.add_page_break()

add_heading("文档修订记录", 1)
add_table(
    ["版本", "日期", "状态", "说明"],
    [["V1.0", "2026-06-18", "产品评审稿", "形成智能经营诊断与整改闭环一期至三期需求方案"]],
    [1200, 1800, 1800, 4560],
    [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading("目录", 1)
for item in [
    "1. 项目背景与问题定义",
    "2. 产品目标与成功指标",
    "3. 产品原则与目标状态",
    "4. 用户角色与使用场景",
    "5. 产品范围与版本规划",
    "6. 总体业务流程",
    "7. 信息架构与页面设计",
    "8. 功能需求",
    "9. 数据需求",
    "10. AI输出规范",
    "11. 权限、安全与审计",
    "12. 异常处理与降级",
    "13. 非功能需求",
    "14. 验收标准与测试场景",
    "15. 排期、依赖与风险",
]:
    add_bullet(item)

doc.add_page_break()

add_heading("1. 项目背景与问题定义", 1)
add_heading("1.1 项目现状", 2)
add_body(
    "现有系统已具备经营数据导入、指标计算、机构排名、横向对比、预警规则、风险评分、经营模式识别和自动报告等能力，覆盖保费、利润、综合成本率、赔付率、费用率及人力效能等核心经营指标。"
)
add_body(
    "当前“AI预警解读”页面的主体分析由本地规则引擎生成，具有确定性和可解释性；项目同时保留大模型流式接口，但尚未在正式前端形成用户可见的调用闭环。"
)

add_heading("1.2 核心问题", 2)
for text in [
    "用户只能阅读固定诊断报告，无法围绕某个结论继续追问或调整分析角度。",
    "报告中的“可能原因”容易被用户理解为已经确认的经营根因。",
    "AI建议不能转化为责任明确、期限明确、目标明确的整改任务。",
    "缺少分险种、分渠道、大额赔案、费用科目等驱动数据，归因深度有限。",
    "缺少证据引用、用户反馈、模型评测、调用审计及权限控制机制。",
]:
    add_bullet(text)

add_heading("1.3 产品机会", 2)
add_callout(
    "核心机会",
    "保留规则引擎作为可信事实底座，让大模型负责解释、追问和表达，再通过整改任务把分析结果转化为经营改善动作。",
)

add_heading("2. 产品目标与成功指标", 1)
add_heading("2.1 总体目标", 2)
for text in [
    "自动发现经营异常，并准确说明异常发生在哪个机构、周期和指标。",
    "在不编造数字的前提下解释异常可能意味着什么。",
    "引导用户核查潜在原因，明确当前数据能够证明和不能证明的内容。",
    "形成可执行的经营改善建议，并支持人工确认后转为整改任务。",
    "跟踪整改进度，在后续数据周期自动复盘改善效果。",
]:
    add_number(text)

add_heading("2.2 成功指标", 2)
add_table(
    ["指标", "一期目标", "测量口径"],
    [
        ("AI数字引用准确率", "≥99%", "AI回答引用的数值与系统计算结果一致"),
        ("无依据结论率", "≤2%", "抽检中缺少数据或规则依据的结论占比"),
        ("AI报告打开率", "≥60%", "打开AI深度解读的报告数/可解读报告数"),
        ("报告追问使用率", "≥30%", "发生至少一次追问的AI报告占比"),
        ("建议采纳或转任务率", "≥20%", "被采纳或转为整改任务的建议占比"),
        ("用户满意率", "≥80%", "用户标记“有帮助”的回答占比"),
        ("AI首字响应时间", "≤5秒", "从提交到前端收到首段流式内容"),
        ("基础报告可用率", "100%", "模型不可用时规则报告仍可访问"),
    ],
    [2880, 1440, 5040],
    [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading("3. 产品原则与目标状态", 1)
add_heading("3.1 三层能力架构", 2)
add_table(
    ["能力层", "主要职责", "约束"],
    [
        ("事实计算层", "指标计算、同比环比、排名、阈值和风险评分", "关键数字不得由大模型自行计算"),
        ("诊断推理层", "组织事实、异常、推断、待核查项和建议", "明确区分事实、规则判断和假设"),
        ("交互生成层", "自然语言解读、追问、摘要改写和任务草稿", "不得越权修改数据、规则或任务状态"),
    ],
    [1800, 3900, 3660],
)
add_callout(
    "核心原则",
    "规则决定“发生了什么”，模型解释“可能意味着什么”，业务人员确认“接下来做什么”。",
)

add_heading("3.2 可信分析标签", 2)
add_table(
    ["标签", "定义", "示例"],
    [
        ("已确认事实", "由当前数据直接计算或读取", "综合成本率159.3%，高于整体59.4个百分点"),
        ("规则判断", "根据已配置阈值或评分规则得出", "触发高风险预警"),
        ("分析推断", "结合经营逻辑形成的可能解释", "风险可能主要来自赔付端"),
        ("待核查", "现有数据不足，需要补充明细", "核查大额赔案和分险种赔付率"),
    ],
    [1800, 3300, 4260],
)

add_heading("4. 用户角色与使用场景", 1)
add_table(
    ["角色", "关注重点", "核心操作"],
    [
        ("总公司管理层", "全国风险、高风险机构、重大异常、整改进展", "查看总览、跨机构比较、下发整改"),
        ("总公司职能部门", "业务、精算、理赔、财务、人力等专业问题", "专业核查、补充原因、处理任务"),
        ("责任区负责人", "责任区机构差距、风险机构、改善进度", "区域比较、追问、创建任务"),
        ("分公司管理层", "本机构风险、短板、核查事项", "查看解读、反馈、处理整改"),
        ("系统管理员", "规则、模型、权限和审计", "配置、监控、审计和问题处理"),
    ],
    [1920, 3840, 3600],
)

add_heading("4.1 典型用户故事", 2)
for text in [
    "作为总公司管理者，我希望快速知道本月哪些分公司风险最高，以及风险主要集中在哪些经营环节。",
    "作为责任区负责人，我希望追问某分公司为何被判定为高风险，并查看每项结论的数据依据。",
    "作为分公司负责人，我希望把系统建议转成整改任务，并明确责任部门、期限和目标。",
    "作为职能部门人员，我希望看到待核查的数据清单，从而判断问题来自赔付、费用、渠道还是人员结构。",
    "作为管理员，我希望追溯每次AI调用所使用的数据、模型和提示词版本。",
]:
    add_bullet(text)

add_heading("5. 产品范围与版本规划", 1)
add_table(
    ["阶段", "产品目标", "主要范围"],
    [
        ("P0 可信AI解读", "完成AI可用闭环", "规则报告、AI深度解读、追问、证据、反馈、降级、日志"),
        ("P1 深度经营分析", "提升归因深度", "分险种/渠道/大案/费用下钻、多机构比较、多月趋势、会议摘要"),
        ("P2 整改闭环", "推动经营改善", "建议转任务、责任与期限、进度提醒、指标复盘、效果评价"),
    ],
    [1800, 2640, 4920],
)

add_heading("5.1 暂不纳入范围", 2)
for text in [
    "AI自动修改经营数据或指标口径。",
    "AI自动启停或调整预警规则。",
    "AI未经用户确认直接创建或下发整改任务。",
    "AI替代核保、理赔、财务或人力审批。",
    "在缺少明细数据时输出确定性经营归因。",
]:
    add_bullet(text)

add_heading("6. 总体业务流程", 1)
for text in [
    "经营数据导入并完成数据质量校验。",
    "系统计算指标、排名、同比环比和机构对标结果。",
    "预警规则识别异常并生成基础诊断报告。",
    "用户按需生成AI深度解读，AI返回结构化分析。",
    "用户围绕当前报告继续追问并查看数据证据。",
    "用户确认建议后生成整改任务草稿。",
    "任务经人工确认后进入执行和跟踪。",
    "进入下一数据周期后，系统自动生成整改复盘。",
]:
    add_number(text)

add_heading("7. 信息架构与页面设计", 1)
add_heading("7.1 导航调整", 2)
add_body("原“AI预警解读（测试）”调整为“智能经营诊断”。")
add_heading("7.2 页面结构", 2)
for text in [
    "全国风险总览：展示高风险、中风险、涉及机构和告警总数。",
    "高风险机构列表：支持按责任区、风险等级、指标类型筛选。",
    "基础诊断报告：展示规则生成的确定性分析。",
    "AI深度解读：按需生成，展示事实、推断、待核查和建议。",
    "交互追问区：提供预置问题和自由提问。",
    "数据证据区：展示指标、周期、对标值和触发规则。",
    "整改任务区：展示任务状态、责任人、期限和复盘结果。",
    "历史记录：查看历史诊断、追问、反馈和整改结果。",
]:
    add_bullet(text)

add_heading("7.3 页面状态信息", 2)
add_body("页面顶部固定展示报告周期、数据更新时间、分析范围、规则报告状态、AI服务状态以及“仅供管理参考”的风险提示。")

add_heading("8. 功能需求", 1)
add_heading("8.1 FR-01 基础诊断报告", 2)
add_body("系统继续通过规则引擎生成确定性报告，作为AI分析的事实底座和服务降级方案。")
for text in [
    "展示风险等级、触发预警、核心指标、机构排名及对标差异。",
    "在存在有效比较期时展示同比或环比。",
    "展示经营模式识别、可能原因、待核查数据和管理建议。",
    "每条结论标记为已确认事实、规则判断、分析推断或待核查。",
    "相同数据和规则必须生成相同报告。",
]:
    add_bullet(text)

add_heading("8.2 FR-02 AI深度解读", 2)
add_body("每份机构报告提供“生成AI深度解读”入口。系统将结构化事实、规则诊断和比较数据传给模型。")
add_table(
    ["输出模块", "内容要求"],
    [
        ("一句话结论", "概括当前机构最值得关注的经营问题"),
        ("关键经营事实", "只引用系统提供的数字，并关联证据"),
        ("主要风险", "说明异常指标及其业务含义"),
        ("潜在驱动因素", "以可能性表述，不得冒充已确认原因"),
        ("待核查数据", "明确下一步需要哪些明细或业务事实"),
        ("管理建议", "按短期、中期或长期组织，可执行、可跟踪"),
        ("数据局限", "说明比较期、粒度或数据完整性限制"),
    ],
    [2400, 6960],
)
add_callout(
    "示例",
    "已确认事实：天津分公司综合成本率159.3%，高于分公司整体59.4个百分点。分析推断：当前偏差主要集中在赔付端，但现有数据无法判断是大额赔案还是业务结构变化导致。建议核查：近3个月大额赔案、分险种赔付率及准备金调整情况。",
)

add_heading("8.3 FR-03 AI交互追问", 2)
add_body("用户可以基于当前机构、当前周期和当前报告继续提问。")
for text in [
    "为什么该分公司被判定为高风险？",
    "哪项指标对综合成本率影响最大？",
    "与分公司整体相比，主要差距在哪里？",
    "管理层本周应优先核查什么？",
    "请生成经营分析会一页摘要。",
    "请将建议整理为整改清单。",
]:
    add_bullet(text)
add_body("会话上下文必须包含当前机构、月份、指标数据、比较数据、预警结果和本轮会话历史；切换机构后必须清空或隔离上下文。")

add_heading("8.4 FR-04 数据证据引用", 2)
add_body("AI核心结论旁提供“查看依据”，展示指标名称、当前值、对标值、差异值、数据周期、数据来源和使用的预警规则。")

add_heading("8.5 FR-05 多机构比较", 2)
add_body("用户可选择2至5家分公司进行比较。AI应回答风险差异、共性问题、标杆机构和差异化改善方向。")
add_body("对于保费、利润等规模型绝对指标，不得脱离机构规模直接判断经营优劣，应同时结合达成率、成本率和业务结构解释。")

add_heading("8.6 FR-06 趋势分析", 2)
add_body("至少存在连续3个月有效数据时，系统方可判断持续改善、持续恶化、高位波动、低位波动或单月异常。数据不足时应明确标记。")

add_heading("8.7 FR-07 整改任务", 2)
add_body("用户可将AI建议转为整改任务草稿。AI只能生成草稿，必须由用户确认后创建。")
add_table(
    ["字段", "来源/要求"],
    [
        ("任务标题", "必填，可由AI生成后编辑"),
        ("来源机构", "系统自动带入"),
        ("来源报告周期", "系统自动带入"),
        ("风险指标", "系统自动带入，可多选"),
        ("问题描述", "AI建议草稿，用户可编辑"),
        ("整改措施", "必填"),
        ("责任部门/责任人", "必填"),
        ("完成期限", "必填"),
        ("当前值/目标值", "当前值自动带入，目标值可配置"),
        ("状态", "待确认、进行中、已完成、已关闭"),
        ("复盘结果", "任务完成或进入下一周期后填写"),
    ],
    [2880, 6480],
)

add_heading("8.8 FR-08 整改复盘", 2)
add_body("进入下一数据周期后，系统比较整改前后指标，输出明显改善、小幅改善、无明显变化、继续恶化或数据不足。")
for text in [
    "展示整改前后数值和变化幅度。",
    "展示与全国、责任区或同类机构对标变化。",
    "判断是否达到用户设定的目标。",
    "提示可能存在的其他影响因素，避免把相关性表述为因果关系。",
]:
    add_bullet(text)

add_heading("8.9 FR-09 用户反馈", 2)
add_body("每次AI回答提供“有帮助”“没有帮助”“数字错误”“结论缺少依据”“建议不可执行”和“其他反馈”。反馈不得自动修改原始数据或预警规则。")

add_heading("9. 数据需求", 1)
add_heading("9.1 一期数据", 2)
for text in [
    "机构、责任区和经营月份。",
    "保费实际、保费计划、经营利润和利润计划。",
    "综合成本率、已赚赔付率和已赚费用率。",
    "人均产能、人均利润、前后台人员和人力成本。",
    "预警规则、风险等级、排名、同比和环比结果。",
]:
    add_bullet(text)

add_heading("9.2 二期建议补充数据", 2)
for text in [
    "分险种保费、赔付率、费用率和综合成本率。",
    "车险/非车险结构、渠道、新单和续保数据。",
    "大额赔案、赔案频度、案均赔款和未决赔款。",
    "准备金调整、固定费用、变动费用及预算偏差。",
    "前台、后台、管理人员结构及岗位产出。",
]:
    add_bullet(text)

add_heading("9.3 AI分析前数据校验", 2)
add_table(
    ["校验项", "处理要求"],
    [
        ("分母为零", "相关比率标记为不可计算，不进入确定性结论"),
        ("单位不一致", "阻止分析并提示统一单位"),
        ("比较期缺失", "不生成同比、环比或持续趋势结论"),
        ("指标口径不一致", "阻止跨期或跨机构比较"),
        ("重复导入", "由数据层提示用户确认覆盖"),
        ("异常极值", "提示人工确认，不自动删除或修正"),
    ],
    [2880, 6480],
)

add_heading("10. AI输出规范", 1)
add_heading("10.1 结构化输出", 2)
add_body("后端应要求模型返回结构化JSON，由前端按字段渲染，禁止将未经处理的模型HTML直接插入页面。")
add_table(
    ["字段", "类型", "说明"],
    [
        ("summary", "string", "一句话经营结论"),
        ("facts", "array", "包含指标、当前值、对标值、周期和证据ID"),
        ("inferences", "array", "推断、置信度及依据"),
        ("investigations", "array", "需要补充核查的数据或业务事实"),
        ("recommendations", "array", "建议、周期、责任角色及关联指标"),
        ("limitations", "array", "数据粒度、完整性和比较限制"),
    ],
    [2400, 1440, 5520],
)

add_heading("10.2 生成约束", 2)
for text in [
    "所有经营数字必须来自系统传入数据。",
    "无比较数据时不得生成同比、环比或行业平均结论。",
    "不得把内部预警阈值描述为监管或行业标准。",
    "推断必须使用“可能”“建议核查”等非确定性表述。",
    "无法判断时明确说明“现有数据不足以判断”。",
    "建议应包含动作对象、执行方向和观察指标。",
]:
    add_bullet(text)

add_heading("11. 权限、安全与审计", 1)
add_heading("11.1 权限矩阵", 2)
add_table(
    ["角色", "数据范围", "AI能力"],
    [
        ("总公司管理层", "全国", "查看、比较、追问、创建任务"),
        ("总公司职能部门", "全国或专业范围", "深度分析、专业核查、任务处理"),
        ("责任区负责人", "本责任区", "查看、追问、创建任务"),
        ("分公司负责人", "本机构", "查看、追问、反馈、处理任务"),
        ("系统管理员", "系统配置范围", "规则、模型、权限和日志管理"),
    ],
    [2280, 3000, 4080],
)

add_heading("11.2 安全要求", 2)
for text in [
    "API密钥只能保存在服务端，不得下发到浏览器。",
    "AI接口必须进行身份验证、数据权限校验和调用限流。",
    "保存模型版本、提示词版本、调用时间、用户、机构和数据周期。",
    "所有AI输出必须进行HTML转义和结构校验。",
    "AI不得执行删除数据、修改配置、修改预警规则等高风险操作。",
    "外部模型使用敏感业务数据前必须完成公司安全与合规评审。",
    "支持关闭生成式AI，仅保留规则诊断。",
]:
    add_bullet(text)

add_heading("12. 异常处理与降级", 1)
add_table(
    ["异常场景", "产品处理"],
    [
        ("AI未配置", "显示基础诊断，隐藏或禁用AI生成入口"),
        ("模型超时", "提示稍后重试，不影响看板和基础报告"),
        ("模型限流", "提示排队或稍后重试"),
        ("返回格式错误", "自动重试一次，仍失败则降级"),
        ("数据不足", "说明缺少的数据，不输出强结论"),
        ("比较期不存在", "隐藏同比、环比和趋势判断"),
        ("用户越权查询", "拒绝请求并记录审计日志"),
        ("模型输出疑似含错误数字", "阻止展示并标记异常，回退基础报告"),
    ],
    [2880, 6480],
)

add_heading("13. 非功能需求", 1)
add_table(
    ["类别", "要求"],
    [
        ("性能", "首字响应≤5秒；完整回答建议≤30秒；支持取消生成"),
        ("稳定性", "AI服务故障不得影响数据看板和规则诊断"),
        ("兼容性", "支持当前看板桌面端和移动端主要浏览器"),
        ("可维护性", "模型、提示词、输出Schema和规则版本可配置"),
        ("可观测性", "记录成功率、延迟、Token消耗、错误类型和降级次数"),
        ("可测试性", "具备固定评测数据集，可回归数字准确性和结论边界"),
        ("可访问性", "关键状态不能仅依赖颜色表达，交互元素可键盘操作"),
    ],
    [2400, 6960],
)

add_heading("14. 验收标准与测试场景", 1)
add_heading("14.1 一期核心验收标准", 2)
acceptance = [
    "有预警数据时，可以正常生成基础诊断报告。",
    "AI服务正常时，可以返回并展示结构化深度解读。",
    "AI服务异常时，基础诊断报告仍正常展示。",
    "AI引用的数字与看板计算结果一致。",
    "无上期数据时，不生成同比、环比或趋势结论。",
    "用户可以围绕当前机构继续追问。",
    "切换机构后，不携带上一机构的会话上下文。",
    "每条重要结论均可查看数据依据。",
    "AI明确区分事实、规则判断、推断和待核查事项。",
    "用户可以评价AI回答，并记录反馈类型。",
    "AI输出不能修改原始数据、规则或权限。",
    "所有调用均可通过审计日志追溯。",
]
for text in acceptance:
    add_number(text)

add_heading("14.2 专项测试场景", 2)
add_table(
    ["场景", "预期结果"],
    [
        ("高赔付率、费用率正常", "识别赔付端风险，建议核查大案和业务结构"),
        ("高费用率、赔付率正常", "识别费用驱动风险，建议核查费用科目"),
        ("保费达成不足但COR良好", "识别规模或收入端问题，不误判为成本恶化"),
        ("人均产能和人均利润均偏低", "识别人力效能问题并提示组织数据核查"),
        ("只有单月数据", "只描述当前表现，不判断持续趋势"),
        ("缺少已赚保费", "不计算相关比率，不生成确定性成本归因"),
        ("同一数据重复生成", "事实和关键结论保持稳定"),
        ("恶意输入或越权提问", "拒绝执行并记录日志"),
    ],
    [3600, 5760],
)

add_heading("15. 排期、依赖与风险", 1)
add_heading("15.1 建议排期", 2)
add_table(
    ["阶段", "周期", "交付内容"],
    [
        ("第一阶段", "2-3周", "页面重命名、AI接口接通、结构化上下文、深度解读、服务降级、可信标签"),
        ("第二阶段", "2周", "追问、证据查看、预置问题、用户反馈、审计日志、基础评测集"),
        ("第三阶段", "3-4周", "整改任务、权限和状态、下一周期复盘、整改统计"),
    ],
    [1800, 1440, 6120],
)

add_heading("15.2 关键依赖", 2)
for text in [
    "统一并确认核心经营指标口径。",
    "确定大模型供应方式、网络环境和数据安全边界。",
    "补充身份认证、机构权限和用户角色数据。",
    "确定整改任务的责任部门、状态流转和提醒机制。",
    "建立用于AI回归测试的脱敏样本数据集。",
]:
    add_bullet(text)

add_heading("15.3 主要风险与应对", 2)
add_table(
    ["风险", "影响", "应对措施"],
    [
        ("模型编造数字", "管理判断失真", "关键数字由系统计算；Schema校验；展示前逐项核对"),
        ("数据粒度不足", "归因过度", "强制区分推断与事实；输出待核查清单"),
        ("外部模型数据风险", "敏感信息泄露", "安全评审、最小化传输、脱敏、权限和审计"),
        ("建议难以落地", "使用率低", "建议结构化为动作、责任角色、期限和观察指标"),
        ("用户过度依赖AI", "忽略专业判断", "持续提示管理参考属性，关键任务必须人工确认"),
    ],
    [2400, 2520, 4440],
)

add_heading("15.4 一期发布门槛", 2)
for text in [
    "规则报告稳定可用，AI故障不影响原看板。",
    "数字引用准确率达到99%以上。",
    "不存在机构数据越权问题。",
    "关键结论具有可查看的数据证据。",
    "完成不少于30组典型经营场景回归测试。",
    "完成高赔付、高费用、保费不足、人力效能不足和数据缺失专项测试。",
]:
    add_bullet(text)

add_callout(
    "最终产品建议",
    "一期产品名称统一为“智能经营诊断”。基础诊断由规则引擎生成，AI深度解读作为增强能力，两者在页面上明确区分。",
    color=GREEN,
    fill="EAF7F0",
)

# Keep tables from splitting header awkwardly and set document metadata.
doc.core_properties.title = "智能经营诊断与整改闭环 PRD"
doc.core_properties.subject = "分公司经营与效能数据看板 AI 能力升级"
doc.core_properties.keywords = "PRD, AI, 经营诊断, 预警, 整改闭环, 财险"
doc.core_properties.author = "产品团队"

doc.save(OUT)
print(OUT.resolve())
