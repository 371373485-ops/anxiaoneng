from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path("智能经营智能体_下一阶段建设_PRD_V1.0.docx")
FONT = "Microsoft YaHei"
NAVY, BLUE, GRAY, BLACK = "17365D", "2E74B5", "667085", "1A1A1A"
LIGHT, PALE, WHITE = "F2F4F7", "EAF2F8", "FFFFFF"
RED, AMBER, GREEN = "9B1C1C", "9A6700", "176B43"


def set_font(run, size=10.5, bold=False, color=BLACK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), FONT)
    rpr.rFonts.set(qn("w:ascii"), "Arial")
    rpr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, bottom=90, start=120, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    node = tcpr.first_child_found_in("w:tcMar")
    if node is None:
        node = OxmlElement("w:tcMar")
        tcpr.append(node)
    for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        child = node.find(qn(f"w:{key}"))
        if child is None:
            child = OxmlElement(f"w:{key}")
            node.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def set_geometry(table, widths):
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    ind = tblpr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tblpr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[i]))
            tcw.set(qn("w:type"), "dxa")
            cell_margins(cell)


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    trpr.append(flag)


def paragraph_rule(paragraph, color="D9E2F3", size=8):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


doc = Document()
section = doc.sections[0]
section.page_width, section.page_height = Inches(8.5), Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, NAVY, 8, 4),
]:
    style = doc.styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True


def add_numbering(num_id, abstract_id, fmt, text):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    nf = OxmlElement("w:numFmt")
    nf.set(qn("w:val"), fmt)
    lvl.append(nf)
    lt = OxmlElement("w:lvlText")
    lt.set(qn("w:val"), text)
    lvl.append(lt)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    aid = OxmlElement("w:abstractNumId")
    aid.set(qn("w:val"), str(abstract_id))
    num.append(aid)
    numbering.append(num)


add_numbering(41, 41, "bullet", "•")
add_numbering(42, 42, "decimal", "%1.")


def list_item(text, ordered=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    ppr = p._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), "42" if ordered else "41")
    numpr.append(ilvl)
    numpr.append(numid)
    ppr.append(numpr)
    set_font(p.add_run(text))
    return p


def heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def para(text, lead=None, color=BLACK):
    p = doc.add_paragraph()
    if lead and text.startswith(lead):
        set_font(p.add_run(lead), bold=True, color=color)
        set_font(p.add_run(text[len(lead):]), color=color)
    else:
        set_font(p.add_run(text), color=color)
    return p


def callout(label, text, fill=PALE, color=NAVY):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_geometry(t, [9360])
    c = t.cell(0, 0)
    shade(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(label + " "), bold=True, color=color)
    set_font(p.add_run(text), color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def table(headers, rows, widths, aligns=None, font_size=8.8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_geometry(t, widths)
    repeat_header(t.rows[0])
    for i, value in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, LIGHT)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(str(value)), 9.2, True, NAVY)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = aligns[i] if aligns else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_font(p.add_run(str(value)), font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


# Running header and footer
hp = section.header.paragraphs[0]
set_font(hp.add_run("产品需求文档  |  智能经营智能体下一阶段建设"), 8.5, True, GRAY)
paragraph_rule(hp)
fp = section.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(fp.add_run("PRD V1.0  |  "), 8.5, color=GRAY)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(28)
p.paragraph_format.space_after = Pt(5)
set_font(p.add_run("产品需求文档（下一阶段）"), 12, True, BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
set_font(p.add_run("智能经营智能体建设"), 25, True, NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(20)
set_font(p.add_run("从“带 AI 的经营分析系统”升级为“可靠、可控、可评测的经营智能体”"), 13, color=GRAY)

for key, value in [
    ("版本", "V1.0"),
    ("文档状态", "产品与研发评审稿"),
    ("当前基线", "规则诊断、AI 解读与追问、证据追溯、整改任务、复盘与审计框架"),
    ("目标阶段", "内部受控试点版 / 智能体工作流 V1"),
    ("建议周期", "8 周"),
    ("编制日期", "2026 年 6 月 19 日"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(key + "："), bold=True, color=NAVY)
    set_font(p.add_run(value))
p = doc.add_paragraph()
paragraph_rule(p, BLUE, 12)
callout(
    "核心判断",
    "下一阶段不以增加更多页面为主，而以建立智能体内核为主：让系统能够围绕经营目标自主规划分析步骤、调用确定性工具、校验证据、识别信息缺口、生成可执行建议，并将结果闭环到整改任务和后续复盘。",
)
doc.add_page_break()

heading("文档修订记录", 1)
table(
    ["版本", "日期", "状态", "说明"],
    [("V1.0", "2026-06-19", "评审稿", "定义智能体下一阶段范围、需求、验收标准、实施步骤与里程碑")],
    [1200, 1800, 1600, 4760],
    [WD_ALIGN_PARAGRAPH.CENTER] * 3 + [WD_ALIGN_PARAGRAPH.LEFT],
)

heading("1. 项目背景与阶段判断", 1)
para(
    "当前系统已完成经营数据导入、指标计算、规则预警、结构化诊断、AI 解读与追问、证据追溯、整改任务、后续复盘、审计和基础权限框架。"
    "但系统目前仍以固定流程和单次模型调用为主，属于“AI 分析助手”，距离成熟经营智能体仍缺少可靠评测、工具自主选择、任务规划、过程纠错、跨周期记忆、主动监测和生产级数据治理。"
)
callout(
    "当前成熟度",
    "智能体目标整体完成度约 55%–60%；AI 功能原型约 85%，受控可用诊断助手约 65%，可靠经营智能体约 55%，生产就绪度约 35%–40%。",
    "FFF4E5",
    AMBER,
)

heading("1.1 已具备能力", 2)
for item in [
    "经营数据、计划数据和指标规则的确定性计算能力。",
    "规则诊断与生成式 AI 解耦，AI 故障时基础诊断可继续工作。",
    "AI 结构化解读、连续追问、证据 ID 与部分数字一致性校验。",
    "整改任务创建、状态流转、后续周期复盘和审计记录。",
    "FastAPI 持久化后端、基础角色与机构访问控制框架。",
]:
    list_item(item)

heading("1.2 核心差距", 2)
table(
    ["差距", "当前表现", "目标状态", "优先级"],
    [
        ("评测门禁", "请求成功或返回 JSON 即可能通过", "逐项验证结构、数字、证据、结论与建议", "P0"),
        ("追问可靠性", "自由文本输出未经过完整事实校验", "与诊断报告共用证据和数字校验器", "P0"),
        ("数据安全", "备份隔离和代理信任边界不足", "机构级隔离、可信网关、最小化外发", "P0"),
        ("智能体内核", "固定单轮调用，无显式计划", "目标分解、工具选择、执行、校验、纠错", "P1"),
        ("长期记忆", "主要依赖单次诊断快照", "机构、周期、任务和反馈的受控记忆", "P1"),
        ("主动性", "用户手动触发", "新数据到达后自动检测并发起建议", "P2"),
    ],
    [1700, 2900, 3200, 1560],
    font_size=8.4,
)

heading("2. 下一阶段目标与边界", 1)
heading("2.1 阶段目标", 2)
for item in [
    "形成可审计的智能体执行链路：目标 → 计划 → 工具调用 → 结果校验 → 输出 → 闭环。",
    "AI 输出必须受经营数据、指标元数据和证据快照约束，不允许凭空生成经营数字。",
    "将 AI 评测由“形式检查”升级为真实质量门禁，支持版本对比和发布阻断。",
    "实现诊断、追问、建议、整改和复盘之间的统一上下文。",
    "完成小范围内部受控试点，并积累可用于迭代的真实业务反馈。",
]:
    list_item(item)

heading("2.2 非目标", 2)
for item in [
    "本阶段不允许智能体自动修改源经营数据、计划数据或指标口径。",
    "本阶段不允许智能体绕过人工确认直接关闭整改任务或发布正式经营结论。",
    "本阶段不建设通用聊天机器人，不回答与经营诊断无关的开放领域问题。",
    "本阶段不以替代经营管理人员为目标，AI 结论定位为决策辅助。",
]:
    list_item(item)

heading("2.3 成功标准", 2)
table(
    ["维度", "目标指标", "发布门槛"],
    [
        ("数字可靠性", "AI 引用数字与证据一致", "准确率 ≥ 99%"),
        ("证据可靠性", "事实、结论、建议均可追溯", "有效证据引用率 100%"),
        ("幻觉控制", "无依据结论占比", "≤ 2%，严重错误为 0"),
        ("降级能力", "模型异常时规则诊断可用", "降级成功率 100%"),
        ("任务质量", "建议可转换为整改任务", "关键字段完整率 ≥ 95%"),
        ("用户价值", "试点用户认为结果有帮助", "有帮助率 ≥ 80%"),
    ],
    [1800, 4300, 3260],
)

heading("3. 目标用户与核心场景", 1)
table(
    ["角色", "主要诉求", "权限边界"],
    [
        ("总部经营管理", "跨机构识别风险、比较差异、下发整改", "可查看授权范围汇总与机构数据"),
        ("区域/分公司管理", "理解本机构问题、形成整改行动", "仅可访问授权机构"),
        ("职能部门", "从专业指标角度定位原因和建议", "按职能和机构范围授权"),
        ("系统管理员", "配置模型、规则、评测和审计", "不默认获得业务数据导出权限"),
    ],
    [1900, 4200, 3260],
)
heading("3.1 核心场景", 2)
for item in [
    "数据导入后，系统识别异常机构并生成规则诊断。",
    "用户提出经营目标，智能体拆解分析任务并调用指标、趋势、对标和证据工具。",
    "智能体发现证据不足时，主动说明缺口并向用户提出最少必要的澄清问题。",
    "智能体生成带证据的结论和可执行建议，经用户确认后形成整改任务。",
    "后续周期数据到达时，系统复盘指标变化，并明确相关性边界，不宣称未经证实的因果关系。",
]:
    list_item(item)

heading("4. 智能体目标工作流", 1)
callout(
    "标准执行链",
    "读取用户目标 → 识别机构与周期 → 生成分析计划 → 选择确定性工具 → 执行并保存中间结果 → 校验证据和数字 → 判断是否需要补充信息 → 生成结论与建议 → 用户确认 → 创建整改任务 → 后续复盘。",
)
table(
    ["阶段", "系统行为", "必须留痕"],
    [
        ("目标理解", "识别问题、机构、周期、指标和输出形式", "原始问题、结构化目标"),
        ("计划生成", "拆分分析步骤并说明所需工具", "计划版本、步骤状态"),
        ("工具执行", "调用指标、趋势、对标、证据、任务等工具", "输入、输出、耗时、错误"),
        ("自检纠错", "检查数字、证据、冲突和信息缺口", "校验结果、重试原因"),
        ("结果生成", "区分事实、推断、限制和建议", "模型、Prompt、证据引用"),
        ("闭环执行", "经确认创建整改任务并跟踪复盘", "确认人、任务和状态历史"),
    ],
    [1700, 4700, 2960],
)

heading("5. 功能需求", 1)
heading("5.1 P0：AI 可靠性与安全门禁", 2)
table(
    ["编号", "需求", "关键规则", "验收标准"],
    [
        ("FR-P0-01", "真实 AI 评测引擎", "校验 Schema、数字、证据、允许/禁止结论和必需建议", "空对象、缺证据、错误数字必须失败"),
        ("FR-P0-02", "统一输出校验器", "诊断解读与追问使用同一验证链", "追问中的无依据数字被拦截或降级"),
        ("FR-P0-03", "模型失败处理", "格式错误重试一次；超时、限流、拒答分类降级", "所有异常均有明确状态和审计"),
        ("FR-P0-04", "提示注入防护", "业务数据仅作为数据上下文，不执行其中指令", "注入测试集严重违规为 0"),
        ("FR-P0-05", "数据最小化", "只发送本次任务所需字段，敏感标识可脱敏", "审计可查看字段范围而非敏感正文"),
        ("FR-P0-06", "机构级数据隔离", "备份、诊断、对话、任务均绑定 orgId", "跨机构读取和覆盖测试全部失败"),
        ("FR-P0-07", "可信代理边界", "应用只接受可信网关注入身份，外部不可直连", "伪造身份头无法获得权限"),
    ],
    [1050, 2300, 3300, 2710],
    font_size=8.0,
)

heading("5.2 P1：智能体规划与工具调用", 2)
table(
    ["编号", "需求", "说明", "验收标准"],
    [
        ("FR-P1-01", "目标解析器", "将自然语言转为机构、周期、指标、任务类型和约束", "标准场景解析准确率 ≥ 95%"),
        ("FR-P1-02", "分析计划器", "生成 1–6 个可执行步骤，支持跳过、失败和重试", "每一步均有状态和工具绑定"),
        ("FR-P1-03", "工具注册中心", "统一注册指标、趋势、对标、证据、诊断、任务和复盘工具", "模型只能调用白名单工具"),
        ("FR-P1-04", "确定性计算优先", "数字、排名、变化率必须由工具计算，不由模型心算", "随机抽样数字一致率 100%"),
        ("FR-P1-05", "信息缺口处理", "证据不足时先提问或说明限制，不强行结论", "缺失数据场景无编造结论"),
        ("FR-P1-06", "执行状态机", "planned/running/waiting_user/failed/completed/cancelled", "中断后可恢复且不重复写入"),
        ("FR-P1-07", "结果合成器", "输出事实、推断、建议、限制和证据", "所有事实与建议满足 Schema"),
    ],
    [1050, 2100, 3500, 2710],
    font_size=8.0,
)

heading("5.3 P1：整改与复盘闭环增强", 2)
for item in [
    "每条建议必须绑定 recommendationId、metricId、改善方向、证据 IDs、建议动作和建议周期。",
    "创建整改任务前必须由用户确认责任部门、责任人、截止日期和目标值。",
    "任务状态继续采用 draft → confirmed → in_progress → completed → closed，禁止跨级流转。",
    "复盘必须比较初始值、当前值、目标距离、对标变化和排名变化。",
    "复盘输出必须明确“指标相关性不等于措施因果性”。",
]:
    list_item(item)

heading("5.4 P1：受控记忆", 2)
table(
    ["记忆类型", "保存内容", "禁止内容", "使用方式"],
    [
        ("会话记忆", "当前问题、计划、工具结果", "API Key、原始敏感凭证", "本次执行上下文"),
        ("机构记忆", "已确认口径、历史诊断和整改状态", "未经授权的其他机构信息", "跨周期连续分析"),
        ("用户偏好", "展示粒度、常用指标、输出格式", "敏感身份推断", "个性化呈现"),
        ("反馈记忆", "有帮助、数字错误、缺证据等标签", "未脱敏自由文本外发", "评测与 Prompt 优化"),
    ],
    [1700, 3000, 2600, 2060],
)

heading("5.5 P2：主动监测与试点运营", 2)
for item in [
    "检测新月份或新版本数据到达，自动运行规则诊断和质量检查。",
    "仅在满足触发条件时生成待审阅提醒，不自动发送正式经营结论。",
    "支持按机构、风险级别和指标配置提醒阈值与静默周期。",
    "提供试点反馈面板，展示错误类型、采纳率、任务转化率和模型成本。",
]:
    list_item(item)

heading("6. AI 输出与工具协议", 1)
heading("6.1 标准输出对象", 2)
table(
    ["对象", "必填字段"],
    [
        ("Fact", "id、text、evidenceIds、metricId、value、unit"),
        ("Inference", "id、text、confidence、evidenceIds、limitations"),
        ("Recommendation", "id、title、action、metricId、direction、evidenceIds、ownerRole、period"),
        ("AgentPlan", "goal、steps、requiredTools、status、missingInputs"),
        ("ToolResult", "toolName、inputHash、output、status、latencyMs、errorType"),
    ],
    [2100, 7260],
)
heading("6.2 工具清单", 2)
table(
    ["工具", "职责", "输出原则"],
    [
        ("get_metric_snapshot", "读取指定机构、周期、指标", "返回原始值、单位、版本和来源"),
        ("calculate_metric", "执行确定性指标计算", "返回公式版本和参与字段"),
        ("compare_trend", "比较跨周期变化", "不推断因果"),
        ("compare_benchmark", "与计划、整体或同类机构比较", "明确对标对象"),
        ("get_evidence", "读取证据快照", "仅返回授权机构数据"),
        ("create_remediation_draft", "生成整改草稿", "必须等待用户确认"),
        ("review_remediation", "计算后续改善结果", "输出限制条件"),
    ],
    [2500, 3400, 3460],
)

heading("7. 数据、安全与合规需求", 1)
for item in [
    "所有业务实体统一绑定 orgId；机构名称仅用于展示，不作为唯一权限键。",
    "生产环境必须使用 proxy 或 token 鉴权；proxy 模式下应用不得直接暴露给终端用户。",
    "模型服务密钥只保存在服务端安全环境变量中，严禁写入前端、日志或文档。",
    "发送给模型的数据采用字段白名单；人员姓名、账号等信息按任务需要脱敏或移除。",
    "审计日志记录谁、何时、针对哪个机构和周期、调用了哪个模型与工具、结果状态和错误类型。",
    "数据库写入采用事务；诊断与证据、任务与状态历史必须原子提交。",
    "开发数据库、备份、日志和评测数据不得进入代码仓库或公开部署包。",
]:
    list_item(item)

heading("8. 评测体系与发布门禁", 1)
heading("8.1 回归测试集", 2)
para("首批建立不少于 60 个场景，后续每次发现真实问题都必须沉淀为固定回归用例。")
table(
    ["场景组", "最低数量", "示例"],
    [
        ("数字与单位", "15", "百分比、万元、负数、零分母、四舍五入"),
        ("指标方向", "10", "越高越好、越低越好、目标区间、中性监测"),
        ("证据不足", "8", "缺数据、冲突数据、版本不一致"),
        ("安全边界", "10", "提示注入、越权机构、敏感信息请求"),
        ("建议质量", "10", "建议与指标方向不一致、无法执行、缺责任角色"),
        ("模型异常", "7", "超时、限流、拒答、非 JSON、截断"),
    ],
    [2200, 1500, 5660],
)
heading("8.2 自动评分规则", 2)
for item in [
    "Schema 成功率：输出字段和类型完全符合版本化 Schema。",
    "数字准确率：所有数字均能在证据或确定性工具结果中找到等价值。",
    "证据有效率：evidenceId 存在、属于当前机构与周期，并支持对应陈述。",
    "无依据结论率：结论不能由给定证据支持的比例。",
    "建议完整率：包含指标、方向、动作、责任角色、周期和证据。",
    "降级成功率：模型不可用时规则结果仍可展示且不丢失业务流程。",
]:
    list_item(item)
callout(
    "强制阻断条件",
    "出现跨机构数据泄露、错误经营数字、伪造证据、严重指标方向错误或关键权限绕过时，无论综合分数多少，版本均不得发布。",
    "FDECEC",
    RED,
)

heading("9. 交互与产品要求", 1)
for item in [
    "用户能够查看智能体当前处于“理解问题、分析中、等待补充、生成结果、已降级”等哪一阶段。",
    "默认向用户展示结论与证据，不强制暴露冗长思维过程；可展开查看计划步骤和工具结果摘要。",
    "事实、推断和建议必须采用不同视觉标签，避免用户将推断误认为事实。",
    "每个关键数字均可点击打开证据抽屉，显示来源、周期、口径、公式版本和对标对象。",
    "降级时明确告知“当前展示规则诊断，生成式 AI 未参与”，不能伪装成 AI 成功。",
    "用户可取消长时间任务；取消后停止后续模型调用，并保留已完成步骤供审计。",
]:
    list_item(item)

heading("10. 实施步骤与里程碑", 1)
table(
    ["阶段", "周期", "主要工作", "退出条件"],
    [
        ("阶段 0：基线冻结", "第 1 周", "冻结 Schema、指标元数据、评测集和发布口径", "基线用例可重复运行"),
        ("阶段 1：P0 修复", "第 1–2 周", "评测门禁、追问校验、权限与数据隔离、事务", "P0 测试全绿，无严重风险"),
        ("阶段 2：智能体内核", "第 3–5 周", "目标解析、计划器、工具注册、状态机、自检纠错", "标准任务可端到端执行"),
        ("阶段 3：闭环增强", "第 5–6 周", "受控记忆、建议转任务、复盘增强", "任务字段完整率达标"),
        ("阶段 4：受控试点", "第 7–8 周", "脱敏真实数据、业务专家验收、反馈迭代", "发布门禁和试点指标达标"),
    ],
    [1800, 1300, 3800, 2460],
    font_size=8.4,
)

heading("10.1 建议研发任务拆分", 2)
table(
    ["工作包", "主要产物", "建议负责人"],
    [
        ("WP1 评测与校验", "验证器、评分器、回归集、发布报告", "AI/后端/测试"),
        ("WP2 智能体编排", "计划器、工具注册、状态机、执行记录", "AI/后端"),
        ("WP3 数据与权限", "orgId 隔离、可信代理、事务、审计", "后端/安全"),
        ("WP4 产品交互", "执行状态、证据展示、确认与降级体验", "产品/前端/设计"),
        ("WP5 业务验收", "专家标注、试点用例、反馈闭环", "经营管理/产品/测试"),
    ],
    [2100, 4500, 2760],
)

heading("11. 验收标准", 1)
heading("11.1 P0 验收", 2)
for item in [
    "空对象、错误数字、无效 evidenceId、缺必需结论和禁止结论均能被评测引擎正确识别。",
    "诊断解读和追问均通过统一验证器；连续两次验证失败后自动降级。",
    "分公司角色无法读取、覆盖或推断其他机构备份、诊断、对话和任务。",
    "生产代理模式下，终端伪造身份头不能获得管理员权限。",
    "AI 关闭、超时、限流和格式错误时，规则诊断与整改流程仍可用。",
]:
    list_item(item)
heading("11.2 智能体工作流验收", 2)
for item in [
    "用户提出复杂经营问题后，系统能够生成可执行计划并选择正确工具。",
    "所有经营数字来自工具结果，计划和执行记录可追溯。",
    "缺少关键数据时，智能体进入 waiting_user，而不是生成确定性结论。",
    "用户确认建议后可创建完整整改任务；未确认前不产生正式任务。",
    "任务中断后可恢复，不重复创建诊断、证据或任务记录。",
]:
    list_item(item)

heading("12. 风险与应对", 1)
table(
    ["风险", "影响", "应对措施"],
    [
        ("模型表现波动", "相同输入输出不稳定", "低温度、结构化输出、确定性工具、版本回归"),
        ("业务口径争议", "结论正确但不被认可", "指标元数据版本化，专家确认口径"),
        ("数据外发风险", "敏感经营信息泄露", "字段白名单、脱敏、服务端密钥、审计"),
        ("智能体过度自主", "未经确认执行动作", "高风险动作强制人工确认"),
        ("评测集过拟合", "测试高分但真实使用差", "保留盲测集并持续吸收真实失败案例"),
        ("成本与延迟", "体验差或调用成本失控", "缓存工具结果、限制步骤数和上下文大小"),
    ],
    [2200, 2800, 4360],
)

heading("13. 上线与运营要求", 1)
for item in [
    "先在少量机构和脱敏数据范围内灰度，不直接面向全公司开放。",
    "每次发布生成模型、Prompt、Schema、工具和评测集版本号。",
    "上线后每日监控失败率、降级率、无依据结论率、Token 成本和平均延迟。",
    "每周由产品、业务专家和测试共同复核负反馈案例并决定是否沉淀回归集。",
    "发现严重数字错误、越权或数据泄露时，立即关闭 AI 开关并回退到规则诊断。",
]:
    list_item(item)

heading("14. 评审决策项", 1)
table(
    ["决策项", "建议方案", "需确认方"],
    [
        ("试点范围", "2–3 个机构、2–3 个月脱敏数据", "业务负责人"),
        ("模型部署", "先使用现有服务，保留可替换模型适配层", "技术负责人"),
        ("身份网关", "生产环境必须经过可信内部网关", "安全/基础设施"),
        ("人工确认点", "任务创建、正式结论发布、任务关闭", "产品/业务"),
        ("发布门槛", "采用本文数字、证据和严重错误指标", "产品/测试/业务"),
    ],
    [2100, 4700, 2560],
)

callout(
    "建议立即启动",
    "首先冻结评测基线，并在第一周完成真实评测门禁、追问统一校验和机构级数据隔离设计。只有 P0 可靠性与安全问题关闭后，才进入智能体计划器和工具编排开发。",
    "E8F5EE",
    GREEN,
)

doc.core_properties.title = "智能经营智能体下一阶段建设 PRD"
doc.core_properties.subject = "可靠经营智能体建设需求与实施计划"
doc.core_properties.author = "项目组"
doc.core_properties.keywords = "经营智能体, AI评测, 工具调用, 整改闭环, PRD"
doc.save(OUT)
print(OUT.resolve())
